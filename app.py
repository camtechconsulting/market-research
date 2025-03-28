from flask import Flask, request, jsonify
from flask_cors import CORS
from docx import Document
from datetime import datetime
from docx.shared import Inches
import openai
import os
import fitz
import re
from PIL import Image
import pytesseract

app = Flask(__name__)
CORS(app)

openai.api_key = os.getenv("OPENAI_API_KEY")

REPORT_FOLDER = os.path.join(app.root_path, 'static', 'reports')
LOGO_PATH = os.path.join(app.root_path, 'static', 'logo.png')
os.makedirs(REPORT_FOLDER, exist_ok=True)

def extract_text_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def extract_text_pdf(file):
    pdf = fitz.open(stream=file.read(), filetype="pdf")
    return "\n".join([page.get_text() for page in pdf])

def extract_text_image(file):
    image = Image.open(file.stream)
    return pytesseract.image_to_string(image)

def extract_text(file_storage):
    filename = file_storage.filename.lower()
    if filename.endswith(".pdf"):
        return extract_text_pdf(file_storage)
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        return extract_text_docx(file_storage)
    elif filename.endswith((".png", ".jpg", ".jpeg")):
        return extract_text_image(file_storage)
    else:
        return ""

def clean_markdown(text):
    text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
    text = text.replace('*', '')
    return text.strip()

def extract_table_data(text):
    table = []
    lines = text.strip().splitlines()
    for line in lines:
        if '|' in line:
            row = [cell.strip() for cell in line.split('|') if cell.strip()]
            if row:
                table.append(row)
    return table if len(table) >= 2 else None

def generate_section(prompt):
    try:
        print("Calling OpenAI API...")
        response = openai.ChatCompletion.create(
            model="gpt-4-0125-preview",
            messages=[
                {"role": "system", "content": "You are a market analyst writing structured research reports for business owners. Include tables in Markdown when relevant."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=1500
        )
        return response['choices'][0]['message']['content']
    except Exception as e:
        print("OpenAI API error:", e)
        return "Error generating this section."

def add_logo(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    if os.path.exists(LOGO_PATH):
        run.add_picture(LOGO_PATH, width=Inches(1.73), height=Inches(0.83))
        paragraph.alignment = 1

@app.route('/')
def home():
    return "Market Research Backend is Running!"

@app.route('/generate', methods=['POST'])
def generate_report():
    doc1 = request.files.get('doc1')
    doc2 = request.files.get('doc2')
    doc3 = request.files.get('doc3')

    context = ""
    if doc1: context += extract_text(doc1) + "\n"
    if doc2: context += extract_text(doc2) + "\n"
    if doc3: context += extract_text(doc3) + "\n"

    if not context.strip():
        return jsonify({'error': 'No valid input provided.'}), 400

    doc = Document()
    add_logo(doc)
    doc.add_heading('Market Research Report', 0)

    sections = [
        ("Executive Summary", "Write a concise executive summary of this market research."),
        ("Industry Overview", "Describe the industry background and macro trends relevant to this business."),
        ("Target Market Analysis", "Analyze the demographics, behaviors, and needs of the target audience. Include a table of segments if available."),
        ("Competitive Landscape", "List the main competitors and compare their market positioning using a table."),
        ("Trends & Opportunities", "Describe current trends and opportunities in the market."),
        ("Market Entry Strategy", "What go-to-market strategy is ideal based on this context?"),
        ("Growth Forecast", "Estimate a 3-year financial and customer growth outlook using a chart-style table."),
        ("Final Observations", "Summarize the feasibility and recommend next steps.")
    ]

    for title, instruction in sections:
        doc.add_heading(title, level=1)
        prompt = f"{instruction}\n\nContext:\n{context}"
        gpt_response = generate_section(prompt)
        table_data = extract_table_data(gpt_response)
        if table_data:
            table = doc.add_table(rows=1, cols=len(table_data[0]))
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            for i, val in enumerate(table_data[0]):
                hdr_cells[i].text = val
            for row_data in table_data[1:]:
                row_cells = table.add_row().cells
                for i, val in enumerate(row_data):
                    if i < len(row_cells):
                        row_cells[i].text = val
        else:
            doc.add_paragraph(clean_markdown(gpt_response))

    filename = f"market_research_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    file_path = os.path.join(REPORT_FOLDER, filename)
    doc.save(file_path)

    return jsonify({'download_url': f'/static/reports/{filename}'})

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
